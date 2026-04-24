import io
import re
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

QUESTION_MARKER = "[Question]"
ANSWER_MARKER   = "[Answer]"
SRL_PREFIX      = "[SRL]"
RNK_PREFIX      = "-RNK-"


def _heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _bold_para(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value) if value else "")


# Section header keywords that warrant a Heading 2 in the Word doc
_SECTION_HEADER_PREFIXES = (
    "summary", "deep dive", "metadata", "flashcards", "key terms", "fun facts",
    "memory engine", "core learning", "objective questions", "analytical questions",
    "mcqs", "very short", "short answer", "medium", "long", "assertion-reason",
    "case study", "test 1", "test 2", "section a", "section b", "section c",
    "section d", "section e",
)


def _is_section_header(line: str) -> bool:
    low = line.lower().strip()
    return any(
        low == prefix or low.startswith(prefix + " ") or low.startswith(prefix + "(") or low.startswith(prefix + ":")
        for prefix in _SECTION_HEADER_PREFIXES
    )


def _render_markdown_section(doc: Document, raw: str):
    if not raw:
        doc.add_paragraph("(No content generated)")
        return

    for line in raw.splitlines():
        stripped = line.rstrip()
        if not stripped:
            continue

        # Strip model-internal SRL/RNK tokens before any marker detection
        if stripped.startswith(SRL_PREFIX):
            stripped = stripped[len(SRL_PREFIX):].lstrip()
        elif stripped.startswith(RNK_PREFIX):
            stripped = stripped[len(RNK_PREFIX):].lstrip()

        if not stripped:
            continue

        # [Question] marker — bold paragraph, or labeled field if "Label: value" follows
        if stripped.startswith(QUESTION_MARKER):
            text = stripped[len(QUESTION_MARKER):].strip()
            m = re.match(r"^([A-Za-z][A-Za-z0-9 /\-]{1,24}):\s*(.*)", text)
            if m:
                _bold_para(doc, m.group(1).strip(), m.group(2).strip())
            else:
                p = doc.add_paragraph()
                p.add_run(text).bold = True

        # [Answer] marker — labeled field if "Label: value" follows, otherwise "Answer:" prefix
        elif stripped.startswith(ANSWER_MARKER):
            text = stripped[len(ANSWER_MARKER):].strip()
            m = re.match(r"^([A-Za-z][A-Za-z0-9 /\-]{1,24}):\s*(.*)", text)
            if m:
                _bold_para(doc, m.group(1).strip(), m.group(2).strip())
            else:
                _bold_para(doc, "Answer", text)

        # MCQ options (A)–(D)
        elif re.match(r"^\([A-D]\)\s", stripped):
            doc.add_paragraph(f"    {stripped}")

        # Assertion-Reason options (i)–(iv)
        elif re.match(r"^\(i{1,3}v?\)\s", stripped):
            doc.add_paragraph(f"    {stripped}")

        # Markdown headings
        elif stripped.startswith("### "):
            _heading(doc, stripped[4:], 3)
        elif stripped.startswith("## "):
            _heading(doc, stripped[3:], 2)
        elif stripped.startswith("# "):
            _heading(doc, stripped[2:], 1)

        # Named section headers (before colon check to avoid "Section A: MCQ" splitting wrong)
        elif _is_section_header(stripped):
            _heading(doc, stripped, 2)

        # "Concept N" sub-headings
        elif re.match(r"^Concept\s+\d+", stripped):
            _heading(doc, stripped, 3)

        # Labeled field: Key: value (short alphabetic label before colon)
        elif ":" in stripped and re.match(r"^[A-Za-z][A-Za-z0-9 /\-]{1,24}:", stripped):
            colon = stripped.index(":")
            label = stripped[:colon].strip()
            value = stripped[colon + 1:].strip()
            _bold_para(doc, label, value)

        else:
            doc.add_paragraph(stripped)


def generate_document(results: dict, class_num: str, subject: str, chapter: str) -> bytes:
    t0 = time.time()
    print(f"[DOCGEN] START")

    for k, v in results.items():
        print(f"[DOCGEN]   worker '{k}' result: {len(str(v))} chars")

    doc = Document()

    title = doc.add_heading(f"{chapter} — {subject} Class {class_num}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    sections = [
        ("core_learning",   "Core Learning"),
        ("memory",          "Memory Engine"),
        ("mcqs",            "Multiple Choice Questions"),
        ("very_short",      "Very Short Answer Questions"),
        ("short_answer",    "Short Answer Questions"),
        ("medium_answer",   "Medium Answer Questions"),
        ("long_answer",     "Long Answer Questions"),
        ("assertion_reason","Assertion-Reason Questions"),
        ("case_studies",    "Case Studies"),
        ("tests",           "Test Papers"),
    ]

    for worker_name, display_name in sections:
        raw = results.get(worker_name, "")
        if not raw:
            continue
        print(f"[DOCGEN] Section: {display_name} (+{time.time()-t0:.1f}s)")
        _heading(doc, display_name, 1)
        _render_markdown_section(doc, str(raw))
        doc.add_paragraph()

    print(f"[DOCGEN] Saving to buffer (+{time.time()-t0:.1f}s)")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = buf.read()
    print(f"[DOCGEN] DONE — {len(result)} bytes, total {time.time()-t0:.1f}s")
    return result
